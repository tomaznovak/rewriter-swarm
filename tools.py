from agency_swarm.tools import BaseTool
from pydantic import Field, BaseModel
import pdfplumber
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io, os, shutil
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch

from openai import OpenAI
from dotenv import load_dotenv

load_dotenv
perplexity_api_key = os.environ.get("PERPLEXITY_API_KEY")

class Search_web_tool(BaseTool, BaseModel):
    """This tool searches the web via. Perplexity API."""

    query: str = Field(
        ...,  description = " It caries the info about what, where, why to search for. Necessary information to run the tool."

    )


    def search_tool(self):
        perplexity_client = OpenAI(
        api_key=perplexity_api_key,
        base_url="https://api.perplexity.ai"
        )

        print("Browsing the web...")

        response = perplexity_client.chat.completions.create(
            model="llama-3.1-sonar-huge-128k-online",
            messages=[
                {"role": "system", "content": "You are the best web searcher in the world. Your experience make you an efective and consice searcher."},
                {"role": "user", "content": self.query}
            ]
        )
        return response.choices[0].message.content

    def run(self):
        search_result = self.search_tool()
        return search_result
    
class Rewrite_Tool(BaseTool, BaseModel):

    source_file_path: str = Field(
        ..., description="Path to the original source file. This file is a base template."
    )
    copy_file_name: str = Field(
        ..., description="Path to the copy file. This file is personalized based on the input user information."
    )
    section_start_marker: str = Field(
        ..., description="The marker which marks the start of the section to be rewritten."
    )
    section_end_marker: str = Field(
        ..., description="The marker which marks the end of the section to be rewritten."
    )
    new_text: str = Field(
        ..., description="The text to be written in the section."
    )

    def run(self):
        """This method performs the following steps:
        1. Creates a copy of the source file.
        2. Modifies the copied file.
        3. Replaces the marked text with the desired text.
        """

        # 1. Create a copy of the source file
        try:
            shutil.copy(self.source_file_path, self.copy_file_name)
            print(f"File copied from '{self.source_file_path}' to '{self.copy_file_name}'.")
        except Exception as e:
            print(f"Error copying file: {e}")
            return f"Error copying file: {e}"

        # 2. Read the copied file
        try:
            with open(self.copy_file_name, 'r') as f:
                content = f.read()
        except Exception as e:
            print(f"Error reading file: {e}")
            return f"Error reading file: {e}"

        # 3. Rewrite the specified section in the file
        start_index = content.find(self.section_start_marker)
        end_index = content.find(self.section_end_marker, start_index)

        if start_index == -1 or end_index == -1:
            print('Specified section markers not found in the file.')
            return 'Specified section markers not found in the file.'

        end_index += len(self.section_end_marker)

        original_section = content[start_index:end_index]

        new_section = f"{self.new_text}"

        updated_content = content.replace(original_section, new_section)

        # Write the updated content back to the file
        try:
            with open(self.copy_file_name, 'w') as f:
                f.write(updated_content)
            return f"Section rewritten in '{self.copy_file_name}'."
        except Exception as e:
            print(f"Error writing to file: {e}")
            return f"Error writing to file: {e}"

class FilePathFinder_Tool(BaseTool, BaseModel):

    file_name: str = Field(
        ..., description="Name of the file to be found."
    )

    def run(self):
        file_path_name = os.path.abspath(self.file_name)
        return file_path_name

class PDFRewrite_Tool(BaseTool, BaseModel):

    input_path: str = Field(
        ..., description="Path to the input PDF file."
    )
    output_file_path: str = Field(
        ..., description="Path to the output PDF file. Most likely just the name of the file."
    )
    modifications: Dict[str, str] = Field(
        ..., description="Dictionary mapping old text to new text. New text is replacing old text."
    )

    def run(self):
    # Read the input PDF
        with pdfplumber.open(self.input_path) as pdf:
            # Create a new PDF
            c = canvas.Canvas(self.output_file_path, pagesize=letter)
            
            for page in pdf.pages:
                # Extract text and text objects from the page
                text = page.extract_text()
                text_objects = page.extract_words(x_tolerance=3, y_tolerance=3, keep_blank_chars=True)
                
                # Apply modifications
                modified_text = text
                for old_text, new_text in self.modifications.items():
                    modified_text = modified_text.replace(old_text, new_text)
                
                # Set up the canvas for the current page
                c.setPageSize((page.width, page.height))
                
                # Draw the modified text objects
                for obj in text_objects:
                    # Apply modifications to the text object
                    obj_text = obj['text']
                    for old_text, new_text in self.modifications.items():
                        if old_text in obj_text:
                            obj_text = obj_text.replace(old_text, new_text)
                    
                    # Set font and font size
                    font_name = obj.get('fontname', 'Helvetica')
                    font_size = obj.get('size', 12)
                    try:
                        c.setFont(font_name, font_size)
                    except:
                        # If the font is not available, use a default font
                        c.setFont('Helvetica', font_size)
                    
                    # Draw the text
                    c.drawString(obj['x0'], page.height - obj['top'], obj_text)
                
                c.showPage()
            
            c.save()
            return "PDF modified successfully"

class DocxRewrite_Tool(BaseTool, BaseModel):

    input_path: str = Field(
        ..., description="Path to the input DOCX file."
    )
    output_file_path: str = Field(
        ..., description="Path to the output DOCX file."
    )
    modifications: Dict[str, str] = Field(
        ..., description="Dictionary holding info about what to change."
    )

    def run(self):
        if not os.path.exists(self.input_path):
            raise FileNotFoundError(f"The file {self.input_path} does not exist.")
    
        _, file_extension = os.path.splitext(self.input_path)
        if file_extension.lower() not in ['.docx', '.doc']:
            raise ValueError("The file must be in .docx or .doc format.")

        # Open the document
        doc = Document(self.input_path)

        # Iterate through paragraphs and replace text
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            modified_text = original_text

            for old_text, new_text in self.modifications.items():
                modified_text = modified_text.replace(old_text, new_text)

            if modified_text != original_text:
                # Store original formatting
                runs = paragraph.runs
                paragraph.clear()
                
                # Replace text while preserving formatting
                new_runs = []
                for run in runs:
                    new_run = paragraph.add_run(run.text)
                    # Copy formatting from original run
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.name = run.font.name or 'Arial'
                    new_run.font.size = run.font.size or Pt(11)
                    new_runs.append(new_run)

                # Apply modifications to the new runs
                for new_run in new_runs:
                    for old_text, new_text in self.modifications.items():
                        if old_text in new_run.text:
                            new_run.text = new_run.text.replace(old_text, new_text)

                # Set paragraph alignment
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Save the modified document
        doc.save(self.output_file_path)
        print(f"Modified document saved as: {self.output_file_path}")

class GPDFFile_Tool(BaseTool, BaseModel):

       # def create_custom_pdf(content: str, style: Dict[str, Any], output_path: str) -> str:
    content: str = Field(
        ..., description="This the content of the pdf file."
    )
    style: Dict[str, Any] = Field(
        ..., description="This is the style of the pdf file."
    )
    output_path: str = Field(
        ..., descrioption="This the path to the file created."
    )
    
    def run(self) -> str:
        """
        Create a custom PDF file based on the given content and style specifications.
    
        Args:
            content (str): The main content of the PDF.
            style (Dict[str, Any]): A dictionary containing style specifications.
            output_path (str): The path where the PDF should be saved.
    
        Returns:
        str: The path to the created PDF file.
        """
        # Create a file-like buffer to receive PDF data
        buffer = io.BytesIO()

        # Set up the document based on style specifications
        page_size = A4 if self.style.get('page_size', 'A4') == 'A4' else letter
        pdf = canvas.Canvas(buffer, pagesize=page_size)

        # Set margins
        margins = self.style.get('margins', {'top': 1, 'bottom': 1, 'left': 1, 'right': 1})
        pdf.translate(margins['left'] * inch, margins['bottom'] * inch)

        # Add a title if specified
        if 'title' in self.style:
            pdf.setFont(self.style.get('title_font', 'Helvetica-Bold'), self.style.get('title_size', 18))
            pdf.drawString(0, page_size[1] - margins['top'] * inch, self.style['title'])

        # Add content with specified formatting
        pdf.setFont(self.style.get('content_font', 'Times-Roman'), self.style.get('content_font_size', 12))
        text_object = pdf.beginText(0, page_size[1] - margins['top'] * inch - 1 * inch)
        for line in self.content.split('\n'):
            text_object.textLine(line)
        pdf.drawText(text_object)

        # Add page numbers if specified
        if self.style.get('page_numbers', False):
            pdf.setFont('Helvetica', 8)
            pdf.drawRightString(page_size[0] - margins['right'] * inch, 0.5 * inch, str(pdf.getPageNumber()))

        # Add footer if specified
        if 'footer' in self.style:
            pdf.setFont('Helvetica', 8)
            pdf.drawCentredString(page_size[0] / 2, 0.5 * inch, self.style['footer'])

        # Save the PDF
        pdf.save()

        # Move to the beginning of the StringIO buffer
        buffer.seek(0)
        with open(self.output_path, 'wb') as f:
            f.write(buffer.getvalue())

        return self.output_path
    

class Delete_File_Tool(BaseTool, BaseModel):

    file_path: str = Field(
        ..., description="This is the path to the file to be deleted."
    )

    def run(self) -> bool:
        try:
            if os.path.isfile(self.file_path):
                os.remove(self.file_path)
                return True
            else:
                print(f"File not found: {self.file_path}")
                return False
        except Exception as e:
            print(f"Error deleting file: {e}")
            return False
        

class PDF_Highlight_Tool(BaseTool, BaseModel):
        pass